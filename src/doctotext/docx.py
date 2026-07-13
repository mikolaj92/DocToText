from __future__ import annotations

import copy
from collections.abc import Iterable
from dataclasses import dataclass, field, replace
from io import BytesIO
from pathlib import Path
from typing import Any, Literal
from zipfile import ZipFile

from docx import Document as PyDocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = qn("w:p")
W_R = qn("w:r")
W_T = qn("w:t")


@dataclass(frozen=True)
class TextSegment:
    """Stable textual segment inside a DOCX.

    container_id examples:
      body:p:0
      header:0:p:0
      table:0:r:0:c:0:p:0   (for table cells)

    paragraph_index is the global order index (body + tables + headers/footers),
    counting ALL paragraphs including empty ones for stable addressing.

    run_indices: indices (within the python-docx paragraph.runs) of runs that
    contributed non-empty text at parse time. Useful for domain anchors.
    """

    id: str
    text: str
    part: str
    index: int
    container_id: str | None = None
    paragraph_index: int | None = None
    run_indices: list[int] | None = None

@dataclass(frozen=True)
class SegmentReplacement:
    """Replacement targeting a segment or a sub-range inside it.

    Offsets are in characters of the segment's text.
    If start_offset and end_offset are None -> whole segment.
    """
    container_id: str | None = None
    id: str | None = None
    text: str = ""
    start_offset: int | None = None
    end_offset: int | None = None
InlineSegmentKind = Literal["text", "opaque"]


@dataclass
class InlineSegment:
    """Canonical mechanical segment for paragraph-level DOCX manipulation.

    This is the single source of truth for run/offset addressing, visible-text
    coordinate math, rPr formatting preservation, and opaque inline content
    (images, tabs, breaks, fields, hyperlinks, pre-existing revisions, ...).

    reviewkit (and other consumers) MUST delegate decomposition, splitting,
    insertion, and range replacement to this representation instead of
    reimplementing the logic.

    - kind="text": editable run text; rpr carries formatting to preserve on split/replace.
    - kind="opaque": non-text inline; element is the original XML to re-emit verbatim;
      text holds the visible contribution (e.g. "\t", "\n", or extracted t text) so
      char offsets stay aligned with parser coordinate systems.
    """

    kind: InlineSegmentKind
    text: str
    rpr: Any | None = None
    element: Any | None = None


def _advances_offset(segment: InlineSegment) -> bool:
    """Whether the segment contributes to visible/offset space (base mechanical view)."""
    return segment.kind in ("text", "opaque")


def _visible_text(segments: list[InlineSegment]) -> str:
    return "".join(segment.text for segment in segments if _advances_offset(segment))


def _visible_len(segments: list[InlineSegment]) -> int:
    return len(_visible_text(segments))


def _copy_segment(segment: InlineSegment, text: str) -> InlineSegment:
    return InlineSegment(
        kind=segment.kind,
        text=text,
        rpr=copy.deepcopy(segment.rpr),
        element=copy.deepcopy(segment.element) if segment.element is not None else None,
    )


def _rpr_at(segments: list[InlineSegment], offset: int) -> Any | None:
    """Return a deepcopy of rpr active at the given visible offset."""
    cursor = 0
    previous: Any | None = None
    for segment in segments:
        if segment.kind != "text":
            if _advances_offset(segment):
                cursor += len(segment.text)
            continue
        next_cursor = cursor + len(segment.text)
        if cursor <= offset <= next_cursor:
            return copy.deepcopy(segment.rpr)
        previous = segment.rpr
        cursor = next_cursor
    return copy.deepcopy(previous)


def _index_at_visible_offset(segments: list[InlineSegment], offset: int) -> int:
    cursor = 0
    for index, segment in enumerate(segments):
        if not _advances_offset(segment):
            continue
        if cursor >= offset:
            return index
        cursor += len(segment.text)
        if cursor >= offset:
            return index + 1
    return len(segments)


def _split_visible_offset(segments: list[InlineSegment], offset: int) -> list[InlineSegment]:
    """Split a text segment at the visible character offset. Pure mechanical."""
    if offset <= 0:
        return segments

    result: list[InlineSegment] = []
    cursor = 0
    split_done = False
    for segment in segments:
        if segment.kind != "text":
            result.append(segment)
            if _advances_offset(segment):
                cursor += len(segment.text)
            continue
        next_cursor = cursor + len(segment.text)
        if not split_done and cursor < offset < next_cursor:
            split_at = offset - cursor
            result.append(_copy_segment(segment, segment.text[:split_at]))
            result.append(_copy_segment(segment, segment.text[split_at:]))
            split_done = True
        else:
            result.append(segment)
        cursor = next_cursor
    return result


def _insert_visible(
    segments: list[InlineSegment], offset: int, insert: InlineSegment
) -> list[InlineSegment]:
    """Insert at visible offset. Pure mechanical."""
    segments = _split_visible_offset(segments, offset)
    index = _index_at_visible_offset(segments, offset)
    return [*segments[:index], insert, *segments[index:]]


def _replace_visible_range(
    segments: list[InlineSegment],
    start: int,
    end: int,
    replacement: list[InlineSegment],
) -> list[InlineSegment]:
    """Replace [start, end) visible range. Pure mechanical."""
    segments = _split_visible_offset(_split_visible_offset(segments, end), start)
    result: list[InlineSegment] = []
    inserted = False
    offset = 0
    for segment in segments:
        next_offset = offset + (len(segment.text) if _advances_offset(segment) else 0)
        if segment.kind == "text" and start <= offset and next_offset <= end:
            if not inserted:
                result.extend(s for s in replacement if s.text)
                inserted = True
            offset = next_offset
            continue
        result.append(segment)
        offset = next_offset
    if not inserted:
        index = _index_at_visible_offset(result, start)
        result[index:index] = [s for s in replacement if s.text]
    return result

def _inline_width(child: Any) -> str:
    """Visible contribution of a non-text inline child (tab, break, etc.)."""
    if child.tag == qn("w:tab"):
        return "\t"
    if child.tag in (qn("w:br"), qn("w:cr")):
        return "\n"
    return ""


def _descendant_visible_text(element: Any) -> str:
    """Visible characters contributed by an opaque subtree (for offset accounting)."""
    parts: list[str] = []
    for node in element.iter():
        if node.tag == qn("w:t") and node.text:
            parts.append(node.text)
        elif node.tag == qn("w:tab"):
            parts.append("\t")
        elif node.tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
    return "".join(parts)


def _wrap_run_child(rpr: Any | None, child: Any) -> Any:
    """Wrap a non-text run child back into a run element, preserving rpr if present.
    Used for opaque preservation.
    """
    run = OxmlElement("w:r")  # type: ignore[name-defined]
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    run.append(copy.deepcopy(child))
    return run


def _run_segments(run: Any) -> list[InlineSegment]:
    """Decompose a single <w:r> into text + opaque segments. Pure mechanical."""
    rpr = run.find(qn("w:rPr"))
    result: list[InlineSegment] = []
    for child in run:
        tag = child.tag
        if tag == qn("w:rPr"):
            continue
        if tag == qn("w:t"):
            if child.text:
                result.append(InlineSegment("text", child.text, copy.deepcopy(rpr)))
            continue
        # Non-text run content (tab, break, drawing, field char, ...)
        # re-wrapped so rpr survives on re-emit.
        result.append(
            InlineSegment(
                "opaque",
                _inline_width(child),
                element=_wrap_run_child(rpr, child),
            )
        )
    return result


def paragraph_to_inline_segments(paragraph: Paragraph) -> list[InlineSegment]:
    """Canonical decomposition of a python-docx Paragraph into ordered InlineSegments.

    This is the single mechanical source for:
    - separating editable text runs from opaque inline content
    - preserving rPr on text runs
    - keeping non-text elements (images, tabs, breaks, fields, hyperlinks, ...)
      as opaque with their visible width contribution for offset math.

    Consumers (especially reviewkit) MUST use this instead of re-walking XML.
    """
    segments: list[InlineSegment] = []
    for child in paragraph._p:
        tag = child.tag
        if tag == qn("w:pPr"):
            continue
        if tag == qn("w:r"):
            segments.extend(_run_segments(child))
            continue
        # Opaque top-level element inside paragraph (e.g. a drawing outside a run, or other).
        segments.append(
            InlineSegment(
                "opaque",
                _descendant_visible_text(child),
                element=copy.deepcopy(child),
            )
        )
    if not segments:
        # Match reviewkit fallback: whole paragraph text as one text segment.
        segments = [InlineSegment("text", paragraph.text)]
    return segments


def rebuild_paragraph_from_inline(paragraph: Paragraph, segments: list[InlineSegment]) -> None:
    """Neutral rebuild: replace paragraph children with the given segments.

    - Preserves existing <w:pPr>.
    - Text segments become <w:r><w:rPr>...</w:rPr><w:t>...</w:t></w:r> (best effort rPr).
    - Opaque segments re-emit their original element.
    - No tracked markup, no comments. Pure mechanical roundtrip for non-review use.

    Review-specific rebuild (with ins/del, revision stamping, comment ranges) stays in
    the review layer (reviewkit) which can use this for base then overlay, or keep its
    own emission for tracked semantics.
    """
    parent = paragraph._p
    # Remove existing non-pPr children
    for child in list(parent):
        if child.tag != qn("w:pPr"):
            parent.remove(child)

    for seg in segments:
        if not seg.text and seg.kind != "opaque":
            continue
        if seg.kind == "text":
            run = OxmlElement("w:r")
            if seg.rpr is not None:
                run.append(copy.deepcopy(seg.rpr))
            t = OxmlElement("w:t")
            if seg.text[:1].isspace() or seg.text[-1:].isspace():
                t.set(qn("xml:space"), "preserve")
            t.text = seg.text
            run.append(t)
            parent.append(run)
        elif seg.kind == "opaque" and seg.element is not None:
            parent.append(copy.deepcopy(seg.element))

@dataclass
class _ParaRef:
    """Internal mapping from our segment to python-docx paragraph + metadata."""

    id: str
    container_id: str
    paragraph_index: int
    paragraph: Paragraph
    part_name: str  # "body", "header:0", "table:0:r:0:c:0", etc.


class DocxDocument:
    """DOCX editing surface backed by python-docx (the proper library for the format).

    - Stable container_id + paragraph_index addressing.
    - Whole segment or offset-based partial replacements (run splitting).
    - Preserves formatting because we operate on runs.
    - Roundtrips via python-docx save.
    """

    def __init__(
        self,
        doc: DocxDocumentType,
        segments: list[TextSegment],
        refs: list[_ParaRef],
    ) -> None:
        self._doc = doc
        self._segments = segments
        self._refs = refs  # index-aligned with segments

    @classmethod
    def open(cls, path: str | Path) -> DocxDocument:
        path = Path(path)
        doc = PyDocxDocument(str(path))
        return cls._from_pydocx(doc)

    @classmethod
    def open_bytes(cls, data: bytes) -> DocxDocument:
        doc = PyDocxDocument(BytesIO(data))
        return cls._from_pydocx(doc)

    @classmethod
    def _from_pydocx(cls, doc: DocxDocumentType) -> DocxDocument:
        segments: list[TextSegment] = []
        refs: list[_ParaRef] = []

        # Global paragraph index counts EVERY paragraph in document order
        # (body, table cells, headers, footers), including empty ones.
        # This matches the contract expected by dike_docs locator and anchors.
        global_paragraph_index = 0
        paragraphs_by_index: dict[int, Paragraph] = {}
        paragraphs_by_container: dict[str, Paragraph] = {}

        def add_paragraphs(paragraphs: list[Paragraph], prefix: str) -> None:
            nonlocal global_paragraph_index
            for local_idx, para in enumerate(paragraphs):
                paragraphs_by_index[global_paragraph_index] = para

                # container_id: body uses global index for stability (matches Dike anchors);
                # other sections use local index within their container.
                if prefix == "body":
                    cid = f"body:p:{global_paragraph_index}"
                else:
                    cid = f"{prefix}:p:{local_idx}"

                paragraphs_by_container[cid] = para

                text = "".join(run.text for run in para.runs)
                run_indices = [ri for ri, run in enumerate(para.runs) if run.text] if para.runs else []

                if text:
                    seg_id = f"s{len(segments)}"
                    segments.append(
                        TextSegment(
                            id=seg_id,
                            text=text,
                            part="word/document.xml" if prefix.startswith("body") or prefix.startswith("table") else f"word/{prefix.split(':')[0]}.xml",
                            index=local_idx,
                            container_id=cid,
                            paragraph_index=global_paragraph_index,
                            run_indices=run_indices,
                        )
                    )
                    refs.append(
                        _ParaRef(
                            id=seg_id,
                            container_id=cid,
                            paragraph_index=global_paragraph_index,
                            paragraph=para,
                            part_name=prefix,
                        )
                    )

                global_paragraph_index += 1

        # Body
        add_paragraphs(list(doc.paragraphs), "body")

        # Tables
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    add_paragraphs(list(cell.paragraphs), f"table:{ti}:r:{ri}:c:{ci}")

        # Headers / Footers
        for si, section in enumerate(doc.sections):
            add_paragraphs(list(section.header.paragraphs), f"header:{si}")
            add_paragraphs(list(section.footer.paragraphs), f"footer:{si}")

        instance = cls(doc=doc, segments=segments, refs=refs)
        instance._paragraphs_by_index = paragraphs_by_index
        instance._paragraphs_by_container = paragraphs_by_container
        return instance
    @property
    def segments(self) -> tuple[TextSegment, ...]:
        return tuple(self._segments)

    @property
    def texts(self) -> list[str]:
        return [s.text for s in self._segments]
    # ------------------------------------------------------------------
    # Structure access (generic DOCX addressing - for Temida adapters)
    # ------------------------------------------------------------------

    def resolve_paragraph(self, container_id: str) -> Paragraph | None:
        """Resolve a python-docx Paragraph by stable container_id.

        container_id examples: "body:p:0", "body:p:17", "header:0:p:0",
        "table:0:r:1:c:2:p:0".
        """
        if not hasattr(self, "_paragraphs_by_container"):
            return None
        return self._paragraphs_by_container.get(container_id)

    def resolve_paragraph_by_index(self, index: int) -> Paragraph | None:
        """Resolve by global paragraph index (counts every paragraph in order,
        including empty ones). Matches dike/posejdon locator contracts.
        """
        if not hasattr(self, "_paragraphs_by_index"):
            return None
        return self._paragraphs_by_index.get(index)

    def get_all_paragraphs(self) -> list[Paragraph]:
        """All paragraphs in document order (body, tables, headers, footers).
        Includes empty paragraphs to keep index stable.
        """
        if not hasattr(self, "_paragraphs_by_index") or not self._paragraphs_by_index:
            return []
        max_i = max(self._paragraphs_by_index.keys())
        return [self._paragraphs_by_index[i] for i in range(max_i + 1) if i in self._paragraphs_by_index]

    def get_inline_segments(self, container_id: str) -> list[InlineSegment]:
        """Return the canonical rich InlineSegment decomposition for one paragraph.

        container_id examples: "body:p:0", "header:0:p:0", table cell variants.
        This is the bridge for review-specific layers to obtain the mechanical view
        (text + opaque with rpr/element) and then use the pure offset functions
        (_split_visible_offset, _insert_visible, _replace_visible_range, etc.)
        without reimplementing paragraph traversal or run decomposition.
        """
        para = self.resolve_paragraph(container_id)
        if para is None:
            return []
        return paragraph_to_inline_segments(para)
    # ------------------------------------------------------------------
    # High-level target application (WriteTarget style)
    # ------------------------------------------------------------------

    def apply_targets(
        self,
        targets: list[dict[str, Any] | SegmentReplacement],
        *,
        strict: bool = False,
    ) -> None:
        """Apply a list of replacement targets.

        Each target can be:
          - SegmentReplacement
          - dict with keys: container_id or id, text, optional start_offset/end_offset
          - object with .container_id, .start_offset, .end_offset, .text (e.g. WriteTarget)

        This is the bridge for ReplacementPlan.write_targets.
        """
        normalized: list[dict[str, Any] | SegmentReplacement] = []
        for t in targets:
            if isinstance(t, (dict, SegmentReplacement)):
                normalized.append(t)
            else:
                # duck-type WriteTarget-like
                d = {
                    "container_id": getattr(t, "container_id", None),
                    "id": getattr(t, "segment_id", None),
                    "text": getattr(t, "text", getattr(t, "replacement_text", "")),
                    "start_offset": getattr(t, "start_offset", None),
                    "end_offset": getattr(t, "end_offset", None),
                }
                normalized.append(d)
        self.apply_replacements(normalized, strict=strict)
    def to_markdown(self) -> str:
        blocks = [f"<!-- doctotext:{s.id} -->\n{s.text}" for s in self._segments]
        return "\n\n".join(blocks)

    def get_indexed_paragraphs(self) -> list[tuple[int, str, Paragraph]]:
        """Return every paragraph in document order with its stable identifiers.

        Returns list of (global_paragraph_index, container_id, python-docx.Paragraph).
        Includes empty paragraphs so that paragraph_index stays in sync with
        dike/posejdon anchor contracts (body + tables + headers/footers).
        This is the canonical source of addressing.
        """
        if not hasattr(self, "_paragraphs_by_index") or not self._paragraphs_by_index:
            return []
        max_i = max(self._paragraphs_by_index.keys())
        out: list[tuple[int, str, Paragraph]] = []
        for i in range(max_i + 1):
            if i not in self._paragraphs_by_index:
                continue
            para = self._paragraphs_by_index[i]
            # find a container_id for it (prefer body: global, else scan)
            cid = f"body:p:{i}"
            if cid not in self._paragraphs_by_container:
                for c, p in self._paragraphs_by_container.items():
                    if p is para:
                        cid = c
                        break
            out.append((i, cid, para))
        return out

    # ------------------------------------------------------------------
    # Placeholder replacement (mechanical, for reinjection)
    # ------------------------------------------------------------------

    def replace_placeholder(
        self,
        container_id: str,
        placeholder: str,
        replacement: str,
    ) -> None:
        """Mechanical: in the *current* text of the paragraph identified by container_id,
        find the first occurrence of placeholder and replace it with replacement.

        This is the non-domain part of reinjection flows.
        Offsets are computed on the live paragraph text; run splitting is handled internally.
        """
        para = self.resolve_paragraph(container_id)
        if para is None:
            raise ValueError(f"no paragraph for container_id {container_id!r}")

        current_text = "".join(r.text for r in para.runs)
        start = current_text.find(placeholder)
        if start < 0:
            raise ValueError(
                f"placeholder {placeholder!r} not found in segment {container_id}"
            )
        end = start + len(placeholder)
        self.apply_targets(
            [
                {
                    "container_id": container_id,
                    "text": replacement,
                    "start_offset": start,
                    "end_offset": end,
                }
            ],
            strict=True,
        )

    # ------------------------------------------------------------------
    # Replacement API (supports full + offset ranges via python-docx)
    # ------------------------------------------------------------------

    def apply_texts(self, texts: Iterable[str], *, strict: bool = False) -> None:
        texts = list(texts)
        if len(texts) != len(self._segments):
            raise ValueError(f"expected {len(self._segments)} segments, got {len(texts)}")
        for i, txt in enumerate(texts):
            self._replace_full_segment(i, txt)

    def apply_replacements(
        self,
        replacements: list[dict[str, Any] | SegmentReplacement],
        *,
        strict: bool = False,
    ) -> None:
        by_container = {r.container_id: i for i, r in enumerate(self._segments)}
        by_id = {r.id: i for i, r in enumerate(self._segments)}

        for rep in replacements:
            if isinstance(rep, SegmentReplacement):
                idx = self._find_index(rep, by_container, by_id, strict)
                if idx is None:
                    continue
                self._apply_to_paragraph(idx, rep.text, rep.start_offset, rep.end_offset)
                continue

            # legacy / dict form
            idx = None
            if "container_id" in rep:
                idx = by_container.get(str(rep["container_id"]))
            elif "id" in rep:
                idx = by_id.get(str(rep["id"]))

            if idx is None:
                if strict:
                    raise ValueError(f"unknown target: {rep.get('container_id') or rep.get('id')}")
                continue

            text = str(rep.get("text", ""))
            start = rep.get("start_offset")
            end = rep.get("end_offset")
            self._apply_to_paragraph(idx, text, start, end)
    def apply_markdown(self, markdown: str, *, strict: bool = True) -> None:
        import re as _re
        by_id = {
            m.group("id"): m.group("text").rstrip("\n")
            for m in _re.finditer(r"<!-- doctotext:(?P<id>s\d+) -->\n(?P<text>.*?)(?=\n<!-- doctotext:s\d+ -->\n|\Z)", markdown, _re.DOTALL)
        }
        if strict:
            expected = {s.id for s in self._segments}
            actual = set(by_id.keys())
            missing = sorted(expected - actual)
            unknown = sorted(actual - expected)
            if missing or unknown:
                raise ValueError(f"markdown marker mismatch; missing={missing} unknown={unknown}")

        for i, seg in enumerate(self._segments):
            if seg.id in by_id:
                self._replace_full_segment(i, by_id[seg.id])

    # ------------------------------------------------------------------
    # Save / bytes
    # ------------------------------------------------------------------

    def save_docx(self, path: str | Path) -> None:
        self._doc.save(str(path))

    def to_bytes(self) -> bytes:
        buf = BytesIO()
        self._doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Internal
    # ------------------------------------------------------------------

    def _find_index(
        self,
        rep: SegmentReplacement,
        by_container: dict[str, int],
        by_id: dict[str, int],
        strict: bool,
    ) -> int | None:
        if rep.container_id:
            idx = by_container.get(str(rep.container_id))
        elif rep.id:
            idx = by_id.get(str(rep.id))
        else:
            idx = None
        if idx is None and strict:
            raise ValueError(f"unknown replacement target: {rep.container_id or rep.id}")
        return idx

    def _replace_full_segment(self, index: int, text: str) -> None:
        ref = self._refs[index]
        para = ref.paragraph
        if not para.runs:
            # create one run
            run = para.add_run(text)
        else:
            # replace first run, clear the rest (preserves some formatting on first run)
            para.runs[0].text = text
            for run in para.runs[1:]:
                run.text = ""
        # update our view
        old = self._segments[index]
        self._segments[index] = replace(old, text=text)

    def _apply_to_paragraph(
        self,
        index: int,
        replacement: str,
        start: int | None,
        end: int | None,
    ) -> None:
        ref = self._refs[index]
        para = ref.paragraph
        full = "".join(r.text for r in para.runs)
        s = 0 if start is None else start
        e = len(full) if end is None else end

        if s == 0 and e == len(full):
            self._replace_full_segment(index, replacement)
            return

        # Use run-range logic (adapted from proven posejdon docx_runs)
        ranges = self._build_run_ranges(para)
        # split boundaries (right first)
        for r in reversed(ranges):
            if r[1] < e < r[2]:
                self._split_run(para, r[0], e - r[1])
                break
        ranges = self._build_run_ranges(para)
        for r in reversed(ranges):
            if r[1] < s < r[2]:
                self._split_run(para, r[0], s - r[1])
                break

        # now replace the affected runs
        ranges = self._build_run_ranges(para)
        affected = [r for r in ranges if r[1] >= s and r[2] <= e and r[2] > r[1]]
        if not affected:
            # fallback
            self._replace_full_segment(index, replacement)
            return

        first = True
        for r in affected:
            run = para.runs[r[0]]
            if first:
                run.text = replacement
                first = False
            else:
                run.text = ""

        # refresh our segment text
        new_text = "".join(r.text for r in para.runs)
        old = self._segments[index]
        self._segments[index] = replace(old, text=new_text)

    def _build_run_ranges(self, paragraph: Paragraph) -> list[tuple[int, int, int]]:
        out: list[tuple[int, int, int]] = []
        cur = 0
        for i, run in enumerate(paragraph.runs):
            ln = len(run.text)
            out.append((i, cur, cur + ln))
            cur += ln
        return out

    def _split_run(self, paragraph: Paragraph, run_index: int, offset: int) -> int:
        if offset <= 0:
            return run_index
        run = paragraph.runs[run_index]
        if offset >= len(run.text):
            return run_index + 1
        left = run.text[:offset]
        right = run.text[offset:]
        run.text = left

        # clone the underlying XML element
        cloned = copy.deepcopy(run._element)
        run._element.addnext(cloned)

        new_run = paragraph.runs[run_index + 1]
        new_run.text = right
        return run_index + 1


# ----------------------------------------------------------------------
# Helper to expose for advanced users if needed
# ----------------------------------------------------------------------

def _paragraph_text(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs)
